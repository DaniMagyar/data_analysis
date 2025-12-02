from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

def create_table_doc():
    doc = Document()
    
    # Add Title
    title_paragraph = doc.add_paragraph()
    run = title_paragraph.add_run("Table 1. Baseline firing properties of amygdala neurons.")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    # Table Data
    headers = ["Region", "Cell Type", "n", "Firing Rate (Hz)", "Burst Index"]
    data = [
        [("LA", True), ("PNs", False), "276", "3.33 ± 3.36", "1.40 ± 5.35"],
        [("", False), ("INs", False), "17", "16.58 ± 6.74", "1.21 ± 0.91"],
        [("BA", True), ("PNs", False), "196", "4.39 ± 4.47", "4.57 ± 10.21"],
        [("", False), ("INs", False), "23", "20.98 ± 7.94", "0.71 ± 0.65"],
        [("AStria", True), ("All", False), "120", "4.99 ± 6.80", "0.54 ± 1.17"],
        [("CeA", True), ("All", False), "50", "2.04 ± 2.96", "0.44 ± 1.02"]
    ]
    
    # Create Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Set Header
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header_text
        # Formatting header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)

    # Fill Rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            cell = row_cells[i]
            
            text_val = ""
            is_bold = False
            
            if isinstance(item, tuple):
                text_val = item[0]
                is_bold = item[1]
            else:
                text_val = item
            
            cell.text = text_val
            
            # Formatting cell
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 0 or i == 1: # Align text columns to left? No, center looks okay for short data, but let's stick to center as it's standard for scientific tables often
                    pass
                
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    if is_bold:
                        run.bold = True
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add Note
    note_paragraph = doc.add_paragraph()
    note_run = note_paragraph.add_run("\nValues presented as Mean ± SD. n indicates number of neurons.")
    note_run.italic = True
    note_run.font.name = 'Arial'
    note_run.font.size = Pt(10)

    # Save
    doc.save("manuscript_table.docx")
    print("Successfully created manuscript_table.docx")

if __name__ == "__main__":
    create_table_doc()
