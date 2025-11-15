from docx import Document
from docx.shared import Pt

# Create main figure legend
doc = Document()

# Title as Heading 1
title = doc.add_heading('Figure 4: Monosynaptic responses to fear conditioning stimuli in LA and AStria', level=1)

# Content paragraph with bolded panel labels
p = doc.add_paragraph()
run = p.add_run('(A) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('LA monosynaptic heatmaps showing z-scored PSTHs for CS (left), US (middle), and CS+US (right) stimuli. Only neurons with detected responses in the monosynaptic window (0-25ms post-stimulus) are shown, sorted by cluster (CS-selective, US-selective, multisensory). Color scale shows z-score normalized to baseline (99th percentile). Black lines separate clusters. ')
run.font.size = Pt(10)

run = p.add_run('(B) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('LA delta peak firing rate (ΔFR) bar charts comparing CS, US, and CS+US responses for each cluster. Stacked vertically: CS-selective (top), US-selective (middle), multisensory (bottom). Bars show mean ± SEM. Statistical comparisons use Wilcoxon signed-rank test (*p<0.05, **p<0.01, ***p<0.001). ')
run.font.size = Pt(10)

run = p.add_run('(C) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('AStria monosynaptic heatmaps, same format as panel A. ')
run.font.size = Pt(10)

run = p.add_run('(D) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('AStria delta peak firing rate bar charts, same format as panel B. ')
run.font.size = Pt(10)

run = p.add_run('(E) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Pie charts showing proportions of CS-selective, US-selective, and multisensory neurons among monosynaptic responders in LA (left) and AStria (right). Percentages indicate fraction of monosynaptic neurons in each category. ')
run.font.size = Pt(10)

run = p.add_run('(F) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Across region comparison of delta peak firing rate (ΔFR) for CS (left), US (middle), and CS+US (right) monosynaptic responders comparing LA vs AStria. Individual data points shown as grey circles with jitter. Bars show mean ± SEM. Statistical comparison using Wilcoxon rank-sum test (*p<0.05, **p<0.01, ***p<0.001, n.s.=not significant). Y-axis scaled to 95th percentile to avoid compression from outliers. ')
run.font.size = Pt(10)

run = p.add_run('(G) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Onset latency comparison for monosynaptic responses. Three panels show CS, US, and CS+US onset latencies (ms) comparing LA vs AStria. Bars show mean ± SEM. Statistical comparison using Wilcoxon rank-sum test (*p<0.05, **p<0.01, ***p<0.001, n.s.=not significant).')
run.font.size = Pt(10)

# Save
doc.save('C:\\Users\\dmagyar\\Documents\\data_analysis\\BAfc_claude_code\\figures\\figure_4\\figure_4.docx')
print('Saved figure_4.docx')

# Create supplementary figure legend
doc_supp = Document()

# Title as Heading 1
title = doc_supp.add_heading('Supplementary Figure 4: Monosynaptic onset latency analysis and statistical comparisons', level=1)

# Content paragraph with bolded panel labels
p = doc_supp.add_paragraph()
run = p.add_run('(A) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Onset latency comparison between selective (CS-selective or US-selective) and multisensory (CS&US) neurons. Six boxplots showing: LA CS-selective vs multisensory (CS latency), LA US-selective vs multisensory (US latency), LA CS&US CS vs US latency, AStria CS-selective vs multisensory (CS latency), AStria US-selective vs multisensory (US latency), AStria CS&US CS vs US latency. Box plots show median (center line), interquartile range (box), and full data range (whiskers). Statistical comparisons use Wilcoxon rank-sum test (*p<0.05, **p<0.01, ***p<0.001, n.s.=not significant). ')
run.font.size = Pt(10)

run = p.add_run('(B) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Permutation test validating regional differences in monosynaptic cluster distributions between LA and AStria. Histogram shows distribution of chi-square statistics from 10,000 permutations (grey bars) with observed chi-square value (red line). P-value indicates probability of observing the actual distribution under the null hypothesis of no regional differences. ')
run.font.size = Pt(10)

run = p.add_run('(C) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Contingency table showing observed monosynaptic neuron counts for each cluster type (CS-selective, US-selective, CS&US) in LA and AStria regions. ')
run.font.size = Pt(10)

run = p.add_run('(D) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Statistical summary showing chi-square statistic, permutation p-value, Cramér\'s V effect size, and significance level (***p<0.001, **p<0.01, *p<0.05, n.s.=not significant). ')
run.font.size = Pt(10)

run = p.add_run('(E) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Kruskal-Wallis tests comparing CS vs US vs CS+US stimuli within each cluster for each region (6 panels total). Row 1: LA CS-sel, LA US-sel, LA CS&US. Row 2: AStria CS-sel, AStria US-sel, AStria CS&US. Each panel shows Kruskal-Wallis p-value with significance level (*p<0.05, **p<0.01, ***p<0.001, n.s.=not significant) and post-hoc pairwise Wilcoxon signed-rank test results (CS-US, CS-Both, US-Both). Tests performed on monosynaptic ΔFR (Hz) metric from bar charts in main figure.')
run.font.size = Pt(10)

# Save
doc_supp.save('C:\\Users\\dmagyar\\Documents\\data_analysis\\BAfc_claude_code\\figures\\figure_4\\figure_4_supplementary.docx')
print('Saved figure_4_supplementary.docx')
