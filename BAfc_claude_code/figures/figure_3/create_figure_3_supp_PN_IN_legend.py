from docx import Document
from docx.shared import Pt

# Create supplementary figure legend for PN vs IN comparison
doc = Document()

# Title as Heading 1
title = doc.add_heading('Supplementary Figure 3 - PN vs IN: Cell type comparison of response profiles in lateral amygdala', level=1)

# Content paragraph with bolded panel labels
p = doc.add_paragraph()
run = p.add_run('(A) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('PN (pyramidal neuron) heatmaps showing z-scored PSTHs for CS (left), US (middle), and CS+US (right) stimuli in LA. Neurons sorted by cluster (CS-selective, US-selective, multisensory, non-responsive, inhibited) based on CS and US responses only. Color scale shows z-score normalized to baseline (99th percentile). Black lines separate clusters. ')
run.font.size = Pt(10)

run = p.add_run('(B) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('PN lineplots showing population average firing rates (Hz) for CS-selective (red), US-selective (blue), and multisensory (purple) neurons. Each cluster displayed separately (stacked rows) for CS, US, and CS+US stimuli. Scalebar: 20 Hz. ')
run.font.size = Pt(10)

run = p.add_run('(C) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('IN (interneuron) heatmaps, same format as panel A. ')
run.font.size = Pt(10)

run = p.add_run('(D) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('IN lineplots, same format as panel B. ')
run.font.size = Pt(10)

run = p.add_run('(E) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Response metrics comparing PN vs IN for each cluster across three stimuli. Six bar chart panels arranged as 2 rows (PN top, IN bottom) × 3 columns (CS-selective left, US-selective middle, multisensory right). Each panel shows delta firing rate (ΔFR) for CS, US, and CS+US stimuli. Bars show mean ± SEM. Statistical comparisons use Wilcoxon signed-rank test (*p<0.05, **p<0.01, ***p<0.001). Color coding matches cluster assignments (red = CS-selective, blue = US-selective, purple = multisensory).')
run.font.size = Pt(10)

# Save
doc.save('C:\\Users\\dmagyar\\Documents\\data_analysis\\BAfc_claude_code\\figures\\figure_3\\figure_3_supp_PN_IN.docx')
print('Saved figure_3_supp_PN_IN.docx')
