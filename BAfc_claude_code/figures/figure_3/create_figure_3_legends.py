from docx import Document
from docx.shared import Pt

# Create main figure legend
doc = Document()

# Title as Heading 1
title = doc.add_heading('Figure 3: Regional comparison of CS-selective, US-selective, and multisensory neurons', level=1)

# Content paragraph with bolded panel labels
p = doc.add_paragraph()
run = p.add_run('(A) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('LA heatmaps showing z-scored PSTHs for CS (left), US (middle), and CS+US (right) stimuli. Neurons sorted by cluster (CS-selective, US-selective, multisensory, non-responsive, inhibited) based on CS and US responses only. Color scale shows z-score normalized to baseline (99th percentile). Black lines separate clusters. ')
run.font.size = Pt(10)

run = p.add_run('(B) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('LA lineplots showing population average firing rates (Hz) for CS-selective (red), US-selective (blue), and multisensory (purple) neurons. Each cluster displayed separately (stacked rows) for CS, US, and CS+US stimuli. Scalebar: 20 Hz. ')
run.font.size = Pt(10)

run = p.add_run('(C) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('LA bar charts showing delta firing rate (ΔFR) comparing CS, US, and CS+US responses within each cluster. Bars show mean ± SEM. Statistical comparisons use Wilcoxon signed-rank test (*p<0.05, **p<0.01, ***p<0.001). ')
run.font.size = Pt(10)

run = p.add_run('(D) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('AStria heatmaps, same format as panel A. ')
run.font.size = Pt(10)

run = p.add_run('(E) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('AStria lineplots, same format as panel B. ')
run.font.size = Pt(10)

run = p.add_run('(F) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('AStria bar charts, same format as panel C. ')
run.font.size = Pt(10)

run = p.add_run('(G) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Pie charts showing proportions of CS-selective, US-selective, and multisensory neurons in LA (left) and AStria (right). Percentages indicate fraction of responsive neurons in each category. Legend shows cluster color coding. ')
run.font.size = Pt(10)

run = p.add_run('(H) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Across region comparison of delta peak firing rate (ΔFR) for CS (left), US (middle), and CS+US (right) responsive neurons comparing LA vs AStria. Individual data points shown as grey circles with jitter. Bars show mean ± SEM. Statistical comparison using Wilcoxon rank-sum test (*p<0.05, **p<0.01, ***p<0.001, n.s.=not significant). Y-axis scaled to 95th percentile to avoid compression from outliers.')
run.font.size = Pt(10)

# Save
doc.save('C:\\Users\\dmagyar\\Documents\\data_analysis\\BAfc_claude_code\\figures\\figure_3\\figure_3.docx')
print('Saved figure_3.docx')

# Create supplementary figure legend
doc_supp = Document()

# Title as Heading 1
title = doc_supp.add_heading('Supplementary Figure 3: Statistical analysis of regional cluster distributions and response metrics', level=1)

# Content paragraph with bolded panel labels
p = doc_supp.add_paragraph()
run = p.add_run('(A) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Permutation test validating regional differences in cluster distributions between LA and AStria. Histogram shows distribution of chi-square statistics from 10,000 permutations (grey bars) with observed chi-square value (red line). P-value indicates probability of observing the actual distribution under the null hypothesis of no regional differences. ')
run.font.size = Pt(10)

run = p.add_run('(B) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Contingency table showing observed neuron counts for each cluster type (CS-selective, US-selective, CS+US) in LA and AStria regions. Title displays observed chi-square statistic. ')
run.font.size = Pt(10)

run = p.add_run('(C) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Statistical summary showing chi-square statistic, permutation p-value, Cramér\'s V effect size, and significance level (***p<0.001, **p<0.01, *p<0.05, n.s.=not significant). ')
run.font.size = Pt(10)

run = p.add_run('(D) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Response metrics comparing CS, US, and CS+US stimuli for CS-selective (row 1), US-selective (row 2), and multisensory (row 3) neurons. Three columns show ΔnSpikes (left), ΔFR Hz (middle), and response length ms (right). Darker bars = LA, lighter bars = AStria. Within-region comparisons use Wilcoxon signed-rank test (*p<0.05, **p<0.01, ***p<0.001). Significance brackets positioned at two levels with fixed spacing. LA and AStria labels shown in top row. ')
run.font.size = Pt(10)

run = p.add_run('(E) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Kruskal-Wallis tests comparing CS vs US vs CS+US stimuli within each cluster for each region (6 panels total). Row 5: LA CS-sel, LA US-sel, LA CS&US. Row 6: AStria CS-sel, AStria US-sel, AStria CS&US. Each panel shows Kruskal-Wallis p-value with significance level (*p<0.05, **p<0.01, ***p<0.001, n.s.=not significant) and post-hoc pairwise Wilcoxon signed-rank test results (CS-US, CS-Both, US-Both). Tests performed on ΔFR (Hz) metric.')
run.font.size = Pt(10)

# Save
doc_supp.save('C:\\Users\\dmagyar\\Documents\\data_analysis\\BAfc_claude_code\\figures\\figure_3\\figure_3_supplementary.docx')
print('Saved figure_3_supplementary.docx')
