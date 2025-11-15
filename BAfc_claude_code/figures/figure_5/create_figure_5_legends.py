from docx import Document
from docx.shared import Pt

# Create main figure legend
doc = Document()

# Title as Heading 1
title = doc.add_heading('Figure 5: Optogenetic manipulation of monosynaptic responses', level=1)

# Content paragraph with bolded panel labels
p = doc.add_paragraph()
run = p.add_run('(A) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Example raster plots showing individual neuron responses to CS or US stimuli without light (top row) and with light (middle row) for four example neurons. Each row represents one trial, with spike times shown as vertical marks. Red vertical line indicates stimulus onset. ')
run.font.size = Pt(10)

run = p.add_run('(B) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Firing rate comparison lineplots for the same example neurons showing average firing rate (Hz) over time for no-light (black) and light (red) conditions. Time axis centered at stimulus onset (0 ms). Each plot shows the peristimulus time histogram comparing baseline and optogenetic manipulation conditions. ')
run.font.size = Pt(10)

run = p.add_run('(C) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Pie charts showing proportions of neurons with enhanced (red) vs non-enhanced (gray) monosynaptic responses during optogenetic manipulation. Four panels show LA CS, LA US, AStria CS, and AStria US. Numbers and percentages indicate fraction of responsive neurons in each category. ')
run.font.size = Pt(10)

run = p.add_run('(D) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Spaghetti plots showing trial-averaged spike counts for individual enhanced neurons comparing no-light (left, black) vs light (right, red) conditions. Each gray line represents one neuron, with population mean shown in bold. Four panels correspond to LA CS, LA US, AStria CS, and AStria US enhanced neurons. Enhancement detected using Wilcoxon signed-rank test (p<0.05) comparing spike counts in the monosynaptic window (12-50ms) between light and no-light conditions.')
run.font.size = Pt(10)

# Save
doc.save('C:\\Users\\dmagyar\\Documents\\data_analysis\\BAfc_claude_code\\figures\\figure_5\\figure_5.docx')
print('Saved figure_5.docx')

# Create supplementary figure legend
doc_supp = Document()

# Title as Heading 1
title = doc_supp.add_heading('Supplementary Figure 5: Light-inhibited neurons during optogenetic manipulation', level=1)

# Content paragraph with bolded panel labels
p = doc_supp.add_paragraph()
run = p.add_run('(A) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Example raster plots from a single light-inhibited neuron (7th in LA heatmap) showing responses to CS and US stimuli. Four columns show CS no-light, CS light, US no-light, and US light conditions. Each row represents one trial with spike times shown as vertical marks. Red vertical line indicates stimulus onset. ')
run.font.size = Pt(10)

run = p.add_run('(B) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('LA heatmaps showing z-scored PSTHs for all light-inhibited neurons sorted by hierarchical clustering. Four panels show CS no-light, CS light, US no-light, and US light conditions. Color scale shows z-score normalized to baseline. Neurons sorted to maximize similarity in response patterns. ')
run.font.size = Pt(10)

run = p.add_run('(C) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('LA population average firing rate lineplots for light-inhibited neurons. Four panels show mean firing rate (Hz) over time for CS no-light (black), CS light (red), US no-light (black), and US light (red) conditions. Time axis centered at stimulus onset (0 ms). Shaded areas represent SEM. ')
run.font.size = Pt(10)

run = p.add_run('(D) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('AStria heatmaps showing z-scored PSTHs for all light-inhibited neurons, same format as panel B. ')
run.font.size = Pt(10)

run = p.add_run('(E) ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('AStria population average firing rate lineplots for light-inhibited neurons, same format as panel C. Light-inhibited neurons identified using Wilcoxon ranksum test comparing recent baseline (−0.5 to 0s) vs earlier baseline (−5 to −0.5s) with criterion of p<0.05 and ≥50% firing rate drop.')
run.font.size = Pt(10)

# Save
doc_supp.save('C:\\Users\\dmagyar\\Documents\\data_analysis\\BAfc_claude_code\\figures\\figure_5\\figure_5_supplementary.docx')
print('Saved figure_5_supplementary.docx')
